from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document

from .schemas import AnalysisResult, GenericSheetResult


def render_universal_report(
    input_path: Path,
    table_theme: str,
    workbook_summary: dict[str, Any],
    sheet_results: list[GenericSheetResult],
    analysis: AnalysisResult,
    mode_reason: str,
) -> str:
    lines: list[str] = []
    lines.append("# 表格核心内容总结")
    lines.append("")
    lines.append("## 总体结论")
    lines.append(analysis.llm_insights)
    lines.append("")
    lines.append("## 文件概览")
    lines.append(
        f"本次分析文件为 `{input_path.name}`，采用 `universal` 通用表格分析模式，"
        f"识别出的主体主题为 `{table_theme}`。{mode_reason}"
    )
    lines.append(
        f"工作簿共包含 {workbook_summary.get('sheet_count', 0)} 个工作表，"
        f"总计 {workbook_summary.get('total_rows', 0)} 行、{workbook_summary.get('total_columns', 0)} 列。"
    )
    if workbook_summary.get("focus_sheet"):
        lines.append(f"从数据量和内容密度看，最值得优先关注的工作表是 `{workbook_summary.get('focus_sheet')}`。")
    if workbook_summary.get("best_quality_sheet") or workbook_summary.get("worst_quality_sheet"):
        lines.append(
            f"数据质量最好的是 `{workbook_summary.get('best_quality_sheet')}`，"
            f"相对更需要清洗的是 `{workbook_summary.get('worst_quality_sheet')}`。"
        )
    lines.append("")
    lines.append("## 关键工作表摘要")

    for sheet in sheet_results:
        lines.extend(_render_sheet_summary(sheet))

    lines.append("## 风险提示")
    if analysis.risk_notes:
        for note in analysis.risk_notes:
            lines.append(f"- {note}")
    else:
        lines.append("- 未识别到额外风险提示。")
    lines.append("")

    lines.append("## 建议")
    if analysis.recommendations:
        for suggestion in analysis.recommendations:
            lines.append(f"- {suggestion}")
    else:
        lines.append("- 当前暂无额外建议。")
    lines.append("")
    return "\n".join(lines)


def write_report(content: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    return output_path


def write_docx_summary(
    input_path: Path,
    table_theme: str,
    workbook_summary: dict[str, Any],
    sheet_results: list[GenericSheetResult],
    analysis: AnalysisResult,
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("表格核心内容总结", 0)

    doc.add_heading("总体结论", level=1)
    doc.add_paragraph(analysis.llm_insights)

    doc.add_heading("文件概览", level=1)
    doc.add_paragraph(
        f"文件名称：{input_path.name}\n"
        f"主体主题：{table_theme}\n"
        f"工作表数量：{workbook_summary.get('sheet_count', 0)}\n"
        f"总行数：{workbook_summary.get('total_rows', 0)}\n"
        f"总列数：{workbook_summary.get('total_columns', 0)}"
    )

    doc.add_heading("关键工作表摘要", level=1)
    for sheet in sheet_results:
        doc.add_heading(sheet.sheet_name, level=2)
        for paragraph in _sheet_summary_paragraphs(sheet):
            doc.add_paragraph(paragraph)

    doc.add_heading("风险提示", level=1)
    if analysis.risk_notes:
        for note in analysis.risk_notes:
            doc.add_paragraph(note, style="List Bullet")
    else:
        doc.add_paragraph("未识别到额外风险提示。")

    doc.add_heading("建议", level=1)
    if analysis.recommendations:
        for suggestion in analysis.recommendations:
            doc.add_paragraph(suggestion, style="List Bullet")
    else:
        doc.add_paragraph("当前暂无额外建议。")

    doc.save(output_path)
    return output_path


def _render_sheet_summary(sheet: GenericSheetResult) -> list[str]:
    lines = [f"### {sheet.sheet_name}"]
    for paragraph in _sheet_summary_paragraphs(sheet):
        lines.append(paragraph)
    lines.append("")
    return lines


def _sheet_summary_paragraphs(sheet: GenericSheetResult) -> list[str]:
    paragraphs: list[str] = []
    paragraphs.append(
        f"该工作表被识别为 `{sheet.table_theme}` 类型，"
        f"共有 {sheet.quality_summary['row_count']} 行、{sheet.quality_summary['column_count']} 列。"
    )

    focus_dimensions = "、".join(sheet.profile.get("focus_dimension_columns", [])[:4]) or "无明显维度列"
    focus_numeric = "、".join(sheet.profile.get("focus_numeric_columns", [])[:4]) or "无明显数值列"
    paragraphs.append(f"重点维度列包括：{focus_dimensions}；重点数值列包括：{focus_numeric}。")

    numeric_summaries = sheet.aggregation.get("numeric_summaries", {})
    if numeric_summaries:
        numeric_texts = []
        for column, summary in list(numeric_summaries.items())[:3]:
            numeric_texts.append(
                f"{column}总量约为 {summary['sum']}，均值约为 {summary['mean']}，最大值约为 {summary['max']}"
            )
        paragraphs.append("主要数值信息：" + "；".join(numeric_texts) + "。")

    grouped_summaries = sheet.aggregation.get("grouped_summaries", [])
    if grouped_summaries:
        top_group_texts = []
        for item in grouped_summaries[:2]:
            top_groups = item.get("top_groups", {})
            if top_groups:
                first_key = next(iter(top_groups))
                top_group_texts.append(
                    f"按{item['dimension']}汇总时，{item['numeric']}最高的是 {first_key}（{top_groups[first_key]}）"
                )
        if top_group_texts:
            paragraphs.append("分组汇总结果显示：" + "；".join(top_group_texts) + "。")

    specialized = sheet.specialized_insights
    if specialized.get("template") == "eval_enhancement":
        best_model = specialized.get("best_model", {}).get("model_name")
        worst_model = specialized.get("worst_model", {}).get("model_name")
        if best_model:
            paragraphs.append(f"专题增强结果表明，综合表现最优的模型是 {best_model}。")
        if worst_model:
            paragraphs.append(f"相对表现较弱的模型是 {worst_model}。")
    elif specialized.get("template") == "business_enhancement":
        top_amount_columns = specialized.get("top_amount_columns", [])
        top_dimension_columns = specialized.get("top_dimension_columns", [])
        if top_amount_columns or top_dimension_columns:
            paragraphs.append(
                f"从业务台账视角看，重点金额列为 {('、'.join(top_amount_columns) or '无')}，"
                f"重点分析维度为 {('、'.join(top_dimension_columns) or '无')}。"
            )
    elif specialized.get("template") == "financial_enhancement":
        largest_columns = specialized.get("financial_signals", {}).get("largest_amount_columns", [])
        if largest_columns:
            brief = "；".join(
                f"{item['column']}累计值约 {item['sum']}" for item in largest_columns[:3]
            )
            paragraphs.append("从金额视角看，主要金额列表现为：" + brief + "。")

    if sheet.warnings:
        warning_brief = "；".join(sheet.warnings[:3])
        paragraphs.append(f"需要注意的问题包括：{warning_brief}。")
    else:
        paragraphs.append("当前未发现明显的数据质量问题。")

    return paragraphs
